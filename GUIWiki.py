#GUIWIKI.py
import wikipedia

# python to docx
from docx import Document
def Wiki(keyword,lang = 'th'):
    wikipedia.set_lang(lang)

    #summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2) #เป็นการเพิ่มข้อมูลเข้าไป

    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')

    
#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox


GUI = Tk()
GUI.title('โปรแกรม wiki by Nick')
GUI.geometry('500x400')

#config
FONT1 = ('Angsana New',18)

#คำอธิบาย
L = ttk.Label(GUI,text = 'ค้นหาบทความ',font = FONT1)
L.pack()


#ช่องค้นหาข้อมูล
v_search = StringVar() #กล่องสำหรับเก็บkeyword
E1 = ttk.Entry(GUI,textvariable = v_search,font = FONT1,width = 40)
E1.pack(pady=10)


#ฟังก์ชั่นการค้นหา
def Search():
    keyword = v_search.get() #.get() คือ ดึงข้อมูลเข้ามา
    try:#ลองค้นหาดูว่าได้ผลลัพท์หรือไม่ หากได้ให้ผ่าน
        language = v_radio.get() # th/en/zh
        Wiki(keyword,language)
        messagebox.showinfo('data is saving','search data complete saving data ')
    except:# หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','please write keyword again')
        
        
    #print(wikipedia.search(keyword))# เป็นการบอกว่าคำที่เราsearchมีกี่คำ
    #result = wikipedia.summary(keyword,sentences = 1)
    #print(result)
    

    
#ปู่มค้นหา
B1 = ttk.Button(GUI,text = 'Search',command = Search)
B1.pack(ipadx=20,ipady=10)#ipadx คือ ขยายภายในปุ่มแนวนอน

# เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady = 10)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text = 'ภาษาไทย',variable = v_radio,value = 'th')
RB2 = ttk.Radiobutton(F1,text = 'อังกฤษ',variable = v_radio,value = 'en')
RB3 = ttk.Radiobutton(F1,text = 'จีน',variable = v_radio,value = 'zh')
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row = 0,column = 0)  #grid(row = ..,column = ..)คือการทำให้อยู่ในแนวนอน
RB2.grid(row = 0,column = 1)
RB3.grid(row = 0,column = 2)





GUI.mainloop()
